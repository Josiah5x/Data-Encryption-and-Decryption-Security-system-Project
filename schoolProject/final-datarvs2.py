
from posixpath import expandvars
import sqlite3
import dropbox
from dropbox.exceptions import AuthError
import pathlib
import pandas as pd
from PyQt5 import QtCore, QtGui, QtWidgets
from PyQt5.QtCore import QDir, QDirIterator, QEvent, Qt, QTimer
from os.path import expanduser
from PyQt5.QtWidgets import QHBoxLayout, QHeaderView, QMainWindow, QMenu, QMessageBox, QVBoxLayout, QWidget, QFileSystemModel, QTreeView,  QFileDialog, QGroupBox, QFrame, QToolTip, QDialog, QPlainTextEdit, QProgressBar, QApplication, QLabel, QPushButton, QLineEdit, QTextEdit, QPlainTextEdit
from PyQt5 import QtGui
from PyQt5.Qt import QStandardItemModel, QStandardItem
from PyQt5.QtGui import QIntValidator, QMovie, QFont, QColor, QPixmap
import os, sys, walls, resource, time, projecticons
from cryptography.fernet import Fernet
from datetime import datetime
import qdarkgraystyle, qdarkstyle, qdarktheme
from PyQt5 import QtCore, QtWidgets, QtWebEngineWidgets
from PyQt5.QtWebEngineWidgets import QWebEngineView 
from Crypto.Signature import pkcs1_15
from Crypto.Hash import SHA256
from Crypto.PublicKey import RSA
import socket
import docx
import PyPDF4




class Ui_MainWindow(QWidget):
    
    def setupUi(self, MainWindow):
        MainWindow.setObjectName("MainWindow")
        MainWindow.resize(1195, 705)
        # MainWindow.setStyleSheet("color: #eee;")
        self.centralwidget = QtWidgets.QWidget(MainWindow)
        self.centralwidget.setObjectName("centralwidget")
        self.frame_3 = QtWidgets.QFrame(self.centralwidget)
        self.frame_3.setGeometry(QtCore.QRect(160, 60, 1679, 661))
        font = QtGui.QFont()
        self.frame_3.setFont(font)
        # self.frame_3.setFrameShape(QtWidgets.QFrame.Box)
        # self.frame_3.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_3.setObjectName("frame_3")
        # self.frame_3.setStyleSheet('background: #455268; color: #fff' )
        

        self.comboBox = QtWidgets.QComboBox(self.frame_3)
        self.comboBox.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.comboBox.setGeometry(QtCore.QRect(1, 8, 840, 31))
        self.comboBox.setObjectName("comboBox")
        # self.comboBox.setStyleSheet('background: rgba(238, 238, 238, 255); color: #333' )

        # self.frame = QtWidgets.QFrame(self.frame_3)
        # self.frame.setGeometry(QtCore.QRect(-35, 5, 220, 699))
        # # self.frame.setFrameShape(QtWidgets.QFrame.Box)
        # self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame.setObjectName("frame")

        self.plainTextEdit = QtWidgets.QPlainTextEdit(self.frame_3)
        self.plainTextEdit.setGeometry(QtCore.QRect(850, 40, 350, 60))
        self.plainTextEdit.setObjectName("plainTextEdit")
        # self.plainTextEdit.setFrameShape(QtWidgets.QFrame.Box)
        # self.plainTextEdit.setFrameShadow(QtWidgets.QFrame.Raised)
        self.plainTextEdit.setReadOnly(True)
        # self.plainTextEdit.setStyleSheet("background-color: #fcfdf7;  border: 1px solid #ccc; border-radius: 8px; color: #000; font-size: 15px;")


        self.plainTextEdit2 = QtWidgets.QPlainTextEdit(self.frame_3)
        self.plainTextEdit2.setGeometry(QtCore.QRect(850, 108, 350, 517))
        self.plainTextEdit2.setObjectName("plainTextEdit")
        # self.plainTextEdit2.setFrameShape(QtWidgets.QFrame.Box)
        # self.plainTextEdit2.setFrameShadow(QtWidgets.QFrame.Raised)
        self.plainTextEdit2.setReadOnly(True)
        # self.plainTextEdit2.setStyleSheet("background-color: #fcfdf7; border: 1px solid #ccc; border-radius: 8px; color: #000;")


        # self.plainTextEdit2 = QtWidgets.QPlainTextEdit(self.frame_3)
        # self.plainTextEdit2.setGeometry(QtCore.QRect(850, 145, 320, 475))
        # self.plainTextEdit2.setObjectName("plainTextEdit")
        # # self.plainTextEdit2.setFrameShape(QtWidgets.QFrame.Box)
        # # self.plainTextEdit2.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.plainTextEdit2.setReadOnly(True)
        # self.plainTextEdit2.setStyleSheet("background-color: #fafbfc; border-radius: 8px solid #ccc; color: #000;")



        


        self.model = QtWidgets.QFileSystemModel()
        # self.model.setRootPath('')
        self.treeView = QtWidgets.QTreeView(self.frame_3)
        # self.treeView.setModel(self.model)
        
        # self.listView.setHeaderHidden(True)

        self.treeView.doubleClicked.connect(self.listItemsView)
        self.treeView.setSortingEnabled(True)
        self.treeView.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.treeView.setGeometry(QtCore.QRect(1, 40, 840, 586))
        
        


        self.taskbar = QtWidgets.QFrame(self.centralwidget)
        # self.taskbar.setGeometry(QtCore.QRect(0, 0, 1180, 40))
        self.taskbar.setGeometry(QtCore.QRect(-5, -5, 2220, 60))
        # self.frame.setFrameShape(QtWidgets.QFrame.Box)
        self.taskbar.setFrameShadow(QtWidgets.QFrame.Raised)
        self.taskbar.setObjectName("taskbar")
        self.taskbar.setStyleSheet("background-color: #1c263c;; color: #333;")


        self.sidebar = QtWidgets.QFrame(self.centralwidget)
        self.sidebar.setGeometry(QtCore.QRect(0, 60, 160, 800))
        # self.frame.setFrameShape(QtWidgets.QFrame.Box)
        self.sidebar.setFrameShadow(QtWidgets.QFrame.Raised)
        self.sidebar.setObjectName("sidebar")                    
        # self.sidebar.setStyleSheet("background-color: #000; color: #fff;")

        self.sidelabel = QtWidgets.QLabel(self.sidebar)
        self.sidelabel. setGeometry(QtCore.QRect(0, 0, 160, 40))
        self.sidelabel.setText("Menu")
        self.sidelabel.setStyleSheet("background-color: #8be400; color: #fff; font-weight: bold;")

        # self.openfolder = QPushButton(self.taskbar)
        # self.openfolder.setGeometry(0, 0, 70, 41)
        # # self.openfolder.setFont(QtGui.QFont('consola', 12))
        # self.openfolder.setObjectName("pushButton")
        # self.openfolder.setToolTip('Click to Encrypt')
        # self.openfolder.setText('Open')

        # self.openfolder.clicked.connect(self.addFiles)
        # self.openfolder.setStyleSheet('QPushButton'
        #                                 '{'
        #                                 'background-color: #1655FF;'
        #                                 'color: #fff;'
        #                                 'border: none;'
        #                                 '}'
        #                                 'QPushButton#pushButton:pressed'
        #                                 '{'
        #                                     'padding-left: 5px;'
        #                                     'padding-top: 5px;'
        #                                     '}'
        #                                         )
        

        # self.btn_createkey = QPushButton(self.taskbar)
        # # self.btn_createkey.setGeometry(15, 180, 30, 30)
        # self.btn_createkey.setGeometry(90, 10, 30, 30)
        # self.btn_createkey.setFont(QtGui.QFont('consola', 12))
        # self.btn_createkey.setObjectName("pushButton")
        # self.btn_createkey.clicked.connect(self.dialogCreateKey)
        # self.btn_createkey.setToolTip('Click to Create Key')
        # self.btn_createkey.setStyleSheet('QPushButton'
        #     '{'
        #     'border-image: url(:/image/icons/key_2_32x32.png);'
        #     'color: #333;'
        #      '}'
        #      'QPushButton::Hover'
        #     '{'
        #     'background: #8be400;'
        #     '}')

        # self.btn_createkey_label = QtWidgets.QLabel(self.taskbar)
        # self.btn_createkey_label.setGeometry(QtCore.QRect(80, 35, 80, 30))
        # self.btn_createkey_label.setStyleSheet(
        #                            "text-align: center;\n"
        #                            "color: #333;"
        #                            "background: none")
        # self.btn_createkey_label.setText("Gen-Key")


        # self.btn_chatapp = QPushButton(self.taskbar)
        # # self.btn_chatapp.setGeometry(15, 230, 30, 30)
        # self.btn_chatapp.setGeometry(160, 10, 30, 30)
        # self.btn_chatapp.setFont(QtGui.QFont('consola', 12))
        # self.btn_chatapp.setObjectName("pushButton")
        # self.btn_chatapp.clicked.connect(self.ChatApps)
        # self.btn_chatapp.setToolTip('Click to Chat App')
        # self.btn_chatapp.setStyleSheet('QPushButton'
        #     '{'
        #     'border-image: url(:/image/icons/comment_32x32.png);'
        #     'color: #333;'
        #      '}'
        #      'QPushButton::Hover'
        #     '{'
        #     'background: #2b2b43;'
        #     '}')

        # self.btn_chatapp_label = QtWidgets.QLabel(self.taskbar)
        # self.btn_chatapp_label.setGeometry(QtCore.QRect(140, 35, 80, 30))
        # self.btn_chatapp_label.setStyleSheet(
        #                            "text-align: center;\n"
        #                            "color: #333;\n"
        #                            "background: none")
        # self.btn_chatapp_label.setText("WhatApp Chat")



        
        # self.btn_edit = QPushButton(self.taskbar)
        # # self.btn_edit.setGeometry(15, 285, 30, 30)
        # self.btn_edit.setGeometry(235, 10, 30, 30)
        # self.btn_edit.setFont(QtGui.QFont('consola', 12))
        # self.btn_edit.setObjectName("pushButton")
        # self.btn_edit.clicked.connect(self.hashinfunc)
        # self.btn_edit.setToolTip('Click to Hashing Function')
        # self.btn_edit.setStyleSheet('QPushButton'
        #     '{'
        #     'border-image: url(:/image/icons/edit_32x32.png);'
        #     'color: #333;'
        #      '}'
        #      'QPushButton::Hover'
        #     '{'
        #     'background: #2b2b43;'
        #     '}')

        # self.btn_edit_label = QtWidgets.QLabel(self.taskbar)
        # self.btn_edit_label.setGeometry(QtCore.QRect(220, 35, 80, 30))
        # self.btn_edit_label.setStyleSheet(
        #                            "text-align: center;\n"
        #                            "color: #333;"
        #                            "background: none")
        # self.btn_edit_label.setText("Hashing MSG")

        # self.btn_database = QPushButton(self.taskbar)
        # # self.btn_database.setGeometry(15, 340, 30, 30)
        # self.btn_database.setGeometry(300, 10, 30, 30)
        # self.btn_database.setFont(QtGui.QFont('consola', 12))
        # self.btn_database.setObjectName("pushButton")
        # self.btn_database.clicked.connect(self.dataBaseQuery)
        # self.btn_database.setToolTip('Click to Database Storage')
        # self.btn_database.setStyleSheet('QPushButton'
        #     '{'
        #     'border-image: url(:/image/icons/database_32x32.png);'
        #     'color: #333;'
        #      '}'
        #      'QPushButton::Hover'
        #     '{'
        #     'background: #2b2b43;'
        #     '}')

        # self.btn_database_label = QtWidgets.QLabel(self.taskbar)
        # self.btn_database_label.setGeometry(QtCore.QRect(295, 35, 80, 30))
        # self.btn_database_label.setStyleSheet(
        #                            "text-align: center;\n"
        #                            "color: #333;"
        #                            "background:none")
        # self.btn_database_label.setText("Database")


        # self.email = QPushButton(self.taskbar)
        # # self.email.setGeometry(15, 340, 30, 30)
        # self.email.setGeometry(360, 10, 90, 30)
        # self.email.setFont(QtGui.QFont('consola', 12))
        # self.email.setObjectName("pushButton")
        # self.email.clicked.connect(self.dataBaseQuery)
        # self.email.setToolTip('Click to Send Mail')
        # self.email.setText("Email")
        # self.email.setStyleSheet('QPushButton'
        #     '{'
        #     # 'border-image: url(:/projecticons/imageicons/005b_19.gif);'
        #     'background: #fc0000;'
        #     'color: #fff;'
        #      '}'
        #      'QPushButton::Hover'
        #     '{'
        #     'background: #ff4b4b;'
        #     '}')

        # self.btn_email_label = QtWidgets.QLabel(self.taskbar)
        # self.btn_email_label.setGeometry(QtCore.QRect(365, 35, 80, 30))
        # self.btn_email_label.setStyleSheet(
        #                            "text-align: center;\n"
        #                            "color: #333;"
        #                            "background:none")
        # self.btn_email_label.setText("Email")

        self.gridLayoutWidget = QtWidgets.QWidget(self.sidebar)
        self.gridLayoutWidget.setGeometry(QtCore.QRect(8, 70, 150, 300))
        self.gridLayoutWidget.setObjectName("gridLayoutWidget")
        self.gridLayout = QtWidgets.QGridLayout(self.gridLayoutWidget)
        self.gridLayout.setContentsMargins(0, 0, 0, 0)
        self.gridLayout.setObjectName("gridLayout")
        self.openfile = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.openfile.setGeometry(QtCore.QRect(2, 96, 150, 30))
        self.openfile.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.openfile.setText("Open")
        self.openfile.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/file_32x32.png"))
        self.openfile.setIconSize(QtCore.QSize(25,25))
        self.openfile.setStyleSheet("QPushButton" 
                                    '{'
                                    'border-left: 20px solid #9cff00;'
                                    ' text-align: left;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )

        self.openfile.clicked.connect(self.browse)
        self.openfile.setObjectName("openfile")
        self.gridLayout.addWidget(self.openfile, 0, 1, 1, 1)


        self.openfolder = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.openfolder.setGeometry(QtCore.QRect(2, 135, 150, 30))
        self.openfolder.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/puzzle_32x32.png"))
        self.openfolder.setIconSize(QtCore.QSize(25,25))
        self.openfolder.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.openfolder.setText("Apply theme")
        self.openfolder.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.openfolder.clicked.connect(self.changetheme)
        self.openfolder.setObjectName("openfolder")
        self.gridLayout.addWidget(self.openfolder, 6, 1, 1, 1)


        self.pushButton_3 = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.pushButton_3.setGeometry(QtCore.QRect(2, 175, 150, 30))
        self.pushButton_3.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/key_32x32.png"))
        self.pushButton_3.setIconSize(QtCore.QSize(25,25))
        self.pushButton_3.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.pushButton_3.setText("Creat key")
        self.pushButton_3.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.pushButton_3.clicked.connect(self.dialogCreateKey)
        self.pushButton_3.setObjectName("pushButton_3")
        self.gridLayout.addWidget(self.pushButton_3, 2, 1, 1, 1)
     
        
        self.pushButton_4 = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.pushButton_4.setGeometry(QtCore.QRect(2, 215, 150, 30))
        self.pushButton_4.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.pushButton_4.clicked.connect(self.hashinfunc)
        self.pushButton_4.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/edit_32x32.png"))
        self.pushButton_4.setIconSize(QtCore.QSize(25,25))
        self.pushButton_4.setText("MSG Sign&Verify")
        self.pushButton_4.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.pushButton_4.setObjectName("pushButton_4")
        self.gridLayout.addWidget(self.pushButton_4, 3, 1, 1, 1)


        self.pushButton_5 = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.pushButton_5.setGeometry(QtCore.QRect(2, 255, 150, 30))
        self.pushButton_5.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.pushButton_5.clicked.connect(self.ChatApps)
        self.pushButton_5.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/comment_32x32.png"))
        self.pushButton_5.setIconSize(QtCore.QSize(25,25))
        self.pushButton_5.setText("Chat")
        self.pushButton_5.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.pushButton_5.setObjectName("pushButton_5")
        self.gridLayout.addWidget(self.pushButton_5, 4, 1, 1, 1)


        self.pushButton_6 = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.pushButton_6.setGeometry(QtCore.QRect(2, 295, 150, 30))
        self.pushButton_6.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.pushButton_6.clicked.connect(self.dataBaseQuery)
        self.pushButton_6.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/database_32x32.png"))
        self.pushButton_6.setIconSize(QtCore.QSize(25,25))
        self.pushButton_6.setText("DataBase")
        self.pushButton_6.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.pushButton_6.setObjectName("pushButton_6")
        self.gridLayout.addWidget(self.pushButton_6, 5, 1, 1, 1)



        self.pushButton_7 = QtWidgets.QPushButton(self.gridLayoutWidget)
        # self.pushButton_7.setGeometry(QtCore.QRect(2, 335, 150, 30))
        self.pushButton_7.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.pushButton_7.clicked.connect(self.emailfunc)
        self.pushButton_7.setIcon(QtGui.QIcon("/home/josiah5x/Documents/schoolProject/icons/mail_32x32.png"))
        self.pushButton_7.setIconSize(QtCore.QSize(25,25))
        self.pushButton_7.setText("Email")
        self.pushButton_7.setStyleSheet("QPushButton" 
                                    '{'
                                    'text-align: left;'
                                    'border-left: 20px solid #9cff00;'
                                    # 'color: #fff;'
                                    'font-weight: bold;'
                                    'font-size: 12px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}'
                                    )
        self.pushButton_7.setObjectName("pushButton_7")
        self.gridLayout.addWidget(self.pushButton_7, 1, 1, 1, 1)

      

        self.frame_7 = QtWidgets.QFrame(self.frame_3)
        self.frame_7.setGeometry(QtCore.QRect(850, 5, 291, 35))
        self.frame_7.setStyleSheet("#frame_7{\n"
                                   "\n"
                                   "background: #8be400;\n"
                                   "font-size: 12px;\n"
                                   "font-weight: bold;\n"
                                   "text-align: center;\n"
                                   "color: #333;\n"
                                   "border-top-left-radius: 25px;\n"
                                   "border-bottom-left-radius: 10px;\n"
                                   "border-top-right-radius: 25px;\n"
                                   "border-bottom-right-radius: 10px;\n"
                                   "}")
        self.frame_7.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_7.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_7.setObjectName("frame_7")
        self.label_3 = QtWidgets.QLabel(self.frame_7)
        self.label_3.setGeometry(QtCore.QRect(100, 10, 121, 21))
        self.label_3.setStyleSheet("font-size: 14px;\n"
                                   "font-weight: bold;\n"
                                   "text-align: center;\n"
                                   "background: #8be400;\n"
                                   "color: #fff;"
                                   )
        self.label_3.setObjectName("label_3")

        

        self.frame_13 = QtWidgets.QFrame(self.frame_7)
        self.frame_13.setGeometry(QtCore.QRect(0, 0, 41, 35))
        self.frame_13.setStyleSheet("border-image: url(:/image/org.libreoffice.LibreOffice-writer.ico);\n"
                                    "background:  #008;")
        self.frame_13.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame_13.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame_13.setObjectName("frame_13")
        
        

        self.label_19 = QtWidgets.QLabel(self.taskbar)
        self.label_19.setGeometry(QtCore.QRect(940, 10, 1090, 21))
        self.label_19.setStyleSheet("font-size: 14px;\n"
                                   "font-weight: bold;\n"
                                   "text-align: center;\n"
                                   "background: none;\n"
                                   "color: #fff;")
        self.label_19.setObjectName("label_3")

        self.label_10 = QtWidgets.QLabel(self.taskbar)
        self.label_10.setGeometry(QtCore.QRect(1040, 25, 481, 18))
        self.label_10.setStyleSheet("background: none; font-size: 12px; font-weight: bold;  color: #ccccff;")
        self.label_10.setObjectName("label_10")

        # self.frame_14 = QtWidgets.QFrame(self.centralwidget)
        # self.frame_14.setGeometry(QtCore.QRect(500, 386, 120, 120))
        # self.frame_14.setStyleSheet("background: none; border-image: url(:/image/PngItem_2666090.png);")
        # self.frame_14.setFrameShape(QtWidgets.QFrame.Box)
        # self.frame_14.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_14.setObjectName("frame_14")

        # self.frame_15 = QtWidgets.QFrame(self.centralwidget)
        # self.frame_15.setGeometry(QtCore.QRect(680, 400, 100, 100))
        # self.frame_15.setStyleSheet("background: none; border-image: url(:/image/371-3715555_nasarawa-state-university-keffi-nsuk.png);")
        # self.frame_15.setFrameShape(QtWidgets.QFrame.StyledPanel)
        # self.frame_15.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_15.setObjectName("frame_15")

        


     
        # self.label = QtWidgets.QLabel(self.taskbar)
        # self.label.setGeometry(QtCore.QRect(700, 2, 800, 35))
        # self.label.setStyleSheet("font-size: 12px;\n"
        #                         #  "font-weight: bold;\n"
        #                          "font-family: Raleway;\n"
        #                          "text-align: center;\n"
        #                          "background: none;\n"
        #                          "color: #333;")
        # self.label.setObjectName("label")


        self.label_animation = QLabel(self.sidebar)
        # self.label_animation.setStyleSheet("background: none; color: #fff; border:none;")
        self.movie = QMovie('64x64.gif')
        self.label_animation.setGeometry(40, 500, 80, 100)
        self.label_animation.setMovie(self.movie)
        # self.label_animation.setStyleSheet("background: #9cff00;")

        timer = QTimer()
        self.startAnimation()
        # timer.singleShot(8000)


   


        self.same = 'This is a Final year project for computer Science Nasarawa state University Keffi @ CHUNTAN LIVINUS JOSIAH'
        
        MainWindow.setCentralWidget(self.centralwidget)
        self.menubar = QtWidgets.QMenuBar(MainWindow)
        self.menubar.setGeometry(QtCore.QRect(0, 0, 985, 25))
        self.menubar.setObjectName("menubar")
        self.menuFile = QtWidgets.QMenu(self.menubar)
        self.menuFile.setObjectName("menuFile")
        self.menu = QtWidgets.QMenu(self.menubar)
        self.menu.setObjectName("menu")
        self.menu_2 = QtWidgets.QMenu(self.menubar)
        self.menu_2.setObjectName("menu_2")
        MainWindow.setMenuBar(self.menubar)
        self.statusbar = QtWidgets.QStatusBar(MainWindow)
        self.statusbar.setObjectName("statusbar")
        MainWindow.setStatusBar(self.statusbar)
        self.actionOpenFiles = QtWidgets.QAction(MainWindow)
        self.actionOpenFiles.setObjectName("actionOpenFiles")
        
        self.actionOpenFiles.triggered.connect(self.browse)
        self.actionOpenFolder = QtWidgets.QAction(MainWindow)
        self.actionOpenFolder.setObjectName("actionOpenFolder")
        self.actionOpenFolder.triggered.connect(self.addFiles)
        self.menuFile.addAction(self.actionOpenFiles)
        self.menuFile.addAction(self.actionOpenFolder)
        self.menubar.addAction(self.menuFile.menuAction())
        self.menubar.addAction(self.menu.menuAction())
        self.menubar.addAction(self.menu_2.menuAction())

        self.retranslateUi(MainWindow)
        QtCore.QMetaObject.connectSlotsByName(MainWindow)

    def retranslateUi(self, MainWindow):
        _translate = QtCore.QCoreApplication.translate
        MainWindow.setWindowTitle(_translate("MainWindow", "Encryption and Decryption Security System"))
        MainWindow.statusBar().showMessage(self.same)
        self.label_3.setText(_translate("MainWindow", "DISPLAY Result"))
        # self.label.setText(_translate("MainWindow", ""))
        self.label_19.setText(_translate("MainWindow", "DATA ENCRYPTION AND DECRYPTION SECURITY SYSTEM"))
        self.label_10.setText(_translate("MainWindow", "Your files safety is very important"))
        self.menuFile.setTitle(_translate("MainWindow", "File"))
        self.menu.setTitle(_translate("MainWindow", "View"))
        self.menu_2.setTitle(_translate("MainWindow", "..."))
        self.actionOpenFiles.setText(_translate("MainWindow", "Open Files"))
        self.actionOpenFolder.setText(_translate("MainWindow", "Open Folder"))
        

       
    def theme(self):
       
        app.setStyleSheet(qdarkstyle.load_stylesheet())
        # app.setStyleSheet(qdarktheme.load_stylesheet("light"))
        # app.setStyleSheet(qdarkgraystyle.load_stylesheet())
        
    def changetheme(self):
        self.win = ThemeForm()
        self.win.show()
    def generatekey(self):
        self.win = CreateForm()
        self.win.show()

    def hashinfunc(self):
        self.HashinFunc = VerifyHashing()
        self.HashinFunc.show()

    def emailfunc(self):
        self.Emailform = EmailForm()
        self.Emailform.show()

    def startAnimation(self):
        self.movie.start()


    def dataBaseQuery(self):
        self.theme = DatabaseForm()
        self.theme.show()

    def addFiles(self):
        try:
            # self.listView.clear()
            self.comboBox.clear()
            # self.frame_15.close()
            # self.frame_14.close()

            
            self.folderChoosen = QtWidgets.QFileDialog.getExistingDirectory(self ,'Open file Folder')
            # self.folderChoosen = QtWidgets.QFileDialog.getOpenFileUrl()
            # self.basenamefile = os.path.basename(self.folderChoosen[0])

            self.basenm = os.path.dirname(self.folderChoosen)
            print(self.basenm)

            self.model = QtWidgets.QFileSystemModel()
            self.model.setRootPath(self.basenm)
            self.treeView.setModel(self.model)
            self.treeView.setRootIndex(self.model.index(self.basenm))
            self.treeView.setObjectName("treetView")
            self.treeView.setAlternatingRowColors(True)
            self.treeView.setAnimated(False)
            self.treeView.setIndentation(0)
            self.treeView.setSortingEnabled(True)
            self.treeView.setAllColumnsShowFocus(True)

            self.plainTextEdit.clear()
            self.comboBox.addItem(self.basenm)
            
        except IndexError:
            pass



    def browse(self, signal):
        try:
            
            self.file_path = QtWidgets.QFileDialog.getOpenFileName()
            self.basenm = os.path.dirname(self.file_path[0])
            self.comboBox.clear()
            self.comboBox.insertItem(0, self.basenm)
            self.basenamefile = os.path.basename(self.file_path[0])
           
            

            self.model = QtWidgets.QFileSystemModel()
            self.model.setRootPath(str(self.basenamefile))
            self.treeView.setModel(self.model)
            self.treeView.setRootIndex(self.model.index(str(self.basenm)))
            self.treeView.setColumnWidth(0, 450)
            self.treeView.setAllColumnsShowFocus(True)
            self.treeView.setStyleSheet('QTreeView'
                                    '{ '
                                        # 'alternate-background-color: #9494f8;'
                                        # 'alternate-color: #333;'
                                        'background: #fcfdf7;'
                                        'color: #333;'
                                        # 'font-size: 14px;'
                                        'font-weight: bold;'
                                        # "border: 1px solid #ddd;"
                                        '}'
                                        'QTreeView::Item:selected'
                                        '{'
                                        'background: #4340e9;'
                                        'color: #fff;'
                                        '}'
                                            )
            self.bb = os.path.join(self.basenm, self.basenamefile)
            print(self.basenm)
            

            
            # self.listItemsView()
            self.readtext()
            # self.readtextpdf()
            try:
                with open(f'{self.bb}', 'rt') as self.myfiles:
                    self.myfiles_path = self.myfiles.read()
                    self.plainTextEdit2.setPlainText(self.myfiles_path)
                    ######################################################
                    self.form = Encryption_Decrypt_Form(self)
                    self.form.path_lineEdit.setText(self.bb)
                    self.form.show()
                    
            except UnicodeDecodeError:

                self.plainTextEdit.setPlainText(self.basenamefile)
                ######################################################
                self.form = Encryption_Decrypt_Form(self)
                self.form.path_lineEdit.setText(self.bb)
                self.form.show()
            
            except PermissionError:
                self.plainTextEdit.setPlainText('Files Needed not Folder')



            
            def readtext(self):
                    
                    lst = os.listdir(self.basenm)
                    filec = self.basenamefile
                    for l in lst:
                        if filec in l:
                            split_up = os.path.splitext(filec)
                            file_name = split_up[0]
                            file_extension = split_up[1]
                            if file_extension == '.pdf':

                                self.view = QtWebEngineWidgets.QWebEngineView(self.frame_3)
                                self.settings = self.view.settings()
                                self.settings.setAttribute(QtWebEngineWidgets.QWebEngineSettings.PluginsEnabled, True)
                                self.view.setGeometry(850, 145, 320, 475)
                                url = QtCore.QUrl.fromLocalFile(self.bb)
                                self.view.load(url)
                                self.view.show()




                                # pdfile = open(f'{self.bb}', 'rb' )
                                # read = PyPDF4.PdfFileReader(pdfile)
                                # pgnum = read.numPages
                                # # for g in range(int(pgnum)):
                                # #     # print(g)
                                # #     d = read.getPage(int(g))
                                # #     # fulltext.append(pages.extractText())
                                # display = read.extractText()
                                # self.plainTextEdit2.setPlainText(display)
                                #     # ret = '\n'.join(fulltext)
                                #     # print(ret)
                                # pdfile.close()
                                print('It is Pdf: ' + file_extension )
                            elif file_extension == '.docx':
                                doc = docx.Document(self.bb)
                                fulltext = []
                                for para in doc.paragraphs:
                                    fulltext.append(para.text)
                                ret = '\n'.join(fulltext)
                                self.plainTextEdit2.setPlainText(ret)
                                # dd= self.readtext(self.bb)
                                # print(dd)
                                print('It is Pdf: ' + file_extension )
                            elif file_extension == '.txt':
                                # self.view.close()
                                # self.view.hide()
                                
                                with open(f'{self.bb}', 'rt') as self.myfiles:
                                    self.myfiles_path = self.myfiles.read()
                                    self.plainTextEdit2.setPlainText(self.myfiles_path)

                                # print('It is Pdf: ' + file_extension )
                                self.view.hide()

                            else:
                                print('Extension None')

            

                    def readtextpdf(self):

                        pdfile = open(f'{self.bb}', 'rb' )
                        # with open(f'{self.bb}', 'rb') as self.myfiles:
                        #         self.myfiles_path = self.myfiles.read()
                        read = PyPDF4.PdfFileReader(pdfile)
                        pgnum = read.numPages
                        for g in range(int(pgnum)):
                            # print(g)
                            d = read.getPage(int(g))
                            # fulltext.append(pages.extractText())
                            display = d.extractText()
                            self.plainTextEdit2.setPlainText(display)
                            # ret = '\n'.join(fulltext)
                            # print(ret)
                        pdfile.close()
            
            # self.treeView.reset()
            # self.treeView.insertItem(0, self.basenamefile)
            # self.plainTextEdit.clear()

        except FileNotFoundError:
                pass
    
  
    

    def listItemsView(self, signal):
        self.plainTextEdit.clear()
        self.plainTextEdit2.clear()
        self.file_path = self.model.filePath(signal)
        self.basenm = os.path.dirname(self.file_path)
        self.basenamefile = os.path.basename(self.file_path)
        

        self.bb = os.path.join(self.basenm, self.basenamefile)
        # self.label.setText(self.basenamefile)
        # self.label.setStyleSheet("background: #0b0f12;")
        
        self.comboBox.insertItem(0, self.file_path)

        self.bb = f'{os.path.join(self.file_path)}'
        
        self.readtext()
        # self.readtextpdf()
        

        try:
            with open(f'{self.bb}', 'rt') as self.myfiles:
                self.myfiles_path = self.myfiles.read()
                self.plainTextEdit2.setPlainText(self.myfiles_path)
                
                
        except UnicodeDecodeError:

            self.plainTextEdit.setPlainText(self.basenamefile)
        
        except PermissionError:
            self.plainTextEdit.setPlainText('Files Needed not Folder')
        
        self.form = Encryption_Decrypt_Form(self)
        self.form.path_lineEdit.setText(self.file_path)

        self.form.show()

    def readtext(self):
        
        lst = os.listdir(self.basenm)
        filec = self.basenamefile
        for l in lst:
            if filec in l:
                split_up = os.path.splitext(filec)
                file_name = split_up[0]
                file_extension = split_up[1]
                if file_extension == '.pdf':

                    self.view = QtWebEngineWidgets.QWebEngineView(self.frame_3)
                    self.settings = self.view.settings()
                    self.settings.setAttribute(QtWebEngineWidgets.QWebEngineSettings.PluginsEnabled, True)
                    self.view.setGeometry(850, 145, 320, 475)
                    url = QtCore.QUrl.fromLocalFile(self.bb)
                    self.view.load(url)
                    self.view.show()




                    # pdfile = open(f'{self.bb}', 'rb' )
                    # read = PyPDF4.PdfFileReader(pdfile)
                    # pgnum = read.numPages
                    # # for g in range(int(pgnum)):
                    # #     # print(g)
                    # #     d = read.getPage(int(g))
                    # #     # fulltext.append(pages.extractText())
                    # display = read.extractText()
                    # self.plainTextEdit2.setPlainText(display)
                    #     # ret = '\n'.join(fulltext)
                    #     # print(ret)
                    # pdfile.close()
                    print('It is Pdf: ' + file_extension )
                elif file_extension == '.docx':
                    doc = docx.Document(self.bb)
                    fulltext = []
                    for para in doc.paragraphs:
                        fulltext.append(para.text)
                    ret = '\n'.join(fulltext)
                    self.plainTextEdit2.setPlainText(ret)
                    # dd= self.readtext(self.bb)
                    # print(dd)
                    print('It is Pdf: ' + file_extension )
                elif file_extension == '.txt':
                    # self.view.close()
                    # self.view.hide()
                    
                    with open(f'{self.bb}', 'rt') as self.myfiles:
                        self.myfiles_path = self.myfiles.read()
                        self.plainTextEdit2.setPlainText(self.myfiles_path)

                    # print('It is Pdf: ' + file_extension )
                    self.view.hide()

                else:
                    print('Extension None')

        

    def readtextpdf(self):

        pdfile = open(f'{self.bb}', 'rb' )
        # with open(f'{self.bb}', 'rb') as self.myfiles:
        #         self.myfiles_path = self.myfiles.read()
        read = PyPDF4.PdfFileReader(pdfile)
        pgnum = read.numPages
        for g in range(int(pgnum)):
            # print(g)
            d = read.getPage(int(g))
            # fulltext.append(pages.extractText())
            display = d.extractText()
            self.plainTextEdit2.setPlainText(display)
            # ret = '\n'.join(fulltext)
            # print(ret)
        pdfile.close()

        
    
    
    def combo(self):
        self.comboBox.insertItem(0, self.file_path)
        
    


    def dialogCreateKey(self):
        
        self.key = CreateForm()
        self.key.show()

    def ChatApps(self):

        self.chat = ChatingApp()
        self.chat.show()

#     def dialogDecrypt(self):

#             self.decrypt = DecryptForm()
#             self.decrypt.show()


class ThemeForm(QWidget):
    def __init__(self):
        super().__init__()

        self.left =200
        self.top =400
        self.width =361
        self.height =150
        self.WindowApp()
        # self.setStyleSheet("background: #eff1f6; color: #eee")
        # self.setWindowModality(Qt.ApplicationModal)
        # self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setFixedSize(self.width, self.height)
        # "background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 rgba(234, 234, 234, 255), stop:0.366972 rgba(238, 238, 238, 255), stop:0.958716 rgba(215, 217, 218, 255));\n"
        
        # self.setStyleSheet("background: #fff;")


    def WindowApp(self):
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle(' App')

        self.frame2 = QFrame(self)
        self.frame2.setGeometry(QtCore.QRect(-5, 0, 380, 200))
        self.frame2.setObjectName("frame")
        # self.frame2.setStyleSheet("background: #f2f7d9")


        self.radioButtonSymetric = QtWidgets.QRadioButton(self.frame2)
        self.radioButtonSymetric.setGeometry(40, 10, 160, 31)
        self.radioButtonSymetric.setText('Light Theme')
        self.radioButtonSymetric.setFont(QtGui.QFont('consola', 10))
        self.radioButtonSymetric.setStyleSheet("color: #8be400;")

        self.radioButtonAsymmetric = QtWidgets.QRadioButton(self.frame2)
        self.radioButtonAsymmetric.setGeometry(200, 10, 161, 31)
        self.radioButtonAsymmetric.setText('Dark Theme')
        self.radioButtonAsymmetric.setFont(QtGui.QFont('consola', 10))
        self.radioButtonAsymmetric.setStyleSheet("color: #8be400;")


        # self.keygenerator = QLineEdit(self.frame2)
        # self.keygenerator.setGeometry(10, 70, 341, 41)
        # self.keygenerator.setFont(QtGui.QFont('consola', 10))
        # self.keygenerator.setAlignment(Qt.AlignLeft)
        # self.keygenerator.setPlaceholderText('Type the of the Key')
        # self.keygenerator.setStyleSheet("background: #fff; color: #333; border: 1px solid #ccc;")

        self.btn_submit = QPushButton(self.frame2)
        self.btn_submit.setGeometry(130, 80, 121, 36)
        # self.btn_submit.setFont(QtGui.QFont('consola', 12))
        self.btn_submit.setObjectName("pushButton")
        self.btn_submit.setText("Apply")
        self.btn_submit.setStyleSheet("background: #8be400; color: #fff; border: 4px solid transparent;")
        # self.btn_submit.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.btn_submit.clicked.connect(self.checkedme)

        
        
        # self.progressbar = QProgressBar(self)
        # self.progressbar.setGeometry(10, 120, 341, 8)

    def checkedme(self):
        
        if self.radioButtonAsymmetric.isChecked():
            # for i in range(1001):
            #     time.sleep(0.05)
            #     self.progressbar.setValue(i)
                app.setStyleSheet(qdarkstyle.load_stylesheet())
                app.setStyle('Fusion')
                MainWindow.setStyleSheet("color: #eee;")
        # app.setStyleSheet(qdarktheme.load_stylesheet("light"))
        # app.setStyleSheet(qdarkgraystyle.load_stylesheet())

        if self.radioButtonSymetric.isChecked():
            # for i in range(1001):
            #     time.sleep(0.01)
            #     self.progressbar.setValue(i)
            app.setStyleSheet(qdarktheme.load_stylesheet("light"))
            MainWindow.setStyleSheet("color: #333;")
            app.setStyle('Fusion')
            # app.setStyleSheet(qdarkstyle.load_stylesheet())
            # app.setStyleSheet(qdarkgraystyle.load_stylesheet())
        


  


class CreateForm(QWidget):
    def __init__(self):
        super().__init__()

        self.left =200
        self.top =400
        self.width =361
        self.height =331
        self.WindowApp()
        # self.setStyleSheet("background: #eff1f6; color: #eee")
        # self.setWindowModality(Qt.ApplicationModal)
        # self.setWindowFlag(QtCore.Qt.FramelessWindowHint)
        self.setFixedSize(self.width, self.height)
        # "background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 rgba(234, 234, 234, 255), stop:0.366972 rgba(238, 238, 238, 255), stop:0.958716 rgba(215, 217, 218, 255));\n"
        
        # self.setStyleSheet("background: #fff;")


    def WindowApp(self):
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle(' App')

        self.frame2 = QFrame(self)
        self.frame2.setGeometry(QtCore.QRect(-5, 0, 380, 200))
        self.frame2.setObjectName("frame")
        # self.frame2.setStyleSheet("background: #f2f7d9")


        self.radioButtonSymetric = QtWidgets.QRadioButton(self.frame2)
        self.radioButtonSymetric.setGeometry(10, 10, 160, 31)
        self.radioButtonSymetric.setText('Symetric Single Key')
        self.radioButtonSymetric.setFont(QtGui.QFont('consola', 10))
        self.radioButtonSymetric.setStyleSheet("color: #8be400;")

        self.radioButtonAsymmetric = QtWidgets.QRadioButton(self.frame2)
        self.radioButtonAsymmetric.setGeometry(195, 10, 161, 31)
        self.radioButtonAsymmetric.setText('Asymmetric Pair Keys')
        self.radioButtonAsymmetric.setFont(QtGui.QFont('consola', 10))
        self.radioButtonAsymmetric.setStyleSheet("color: #8be400;")


        self.keygenerator = QLineEdit(self.frame2)
        self.keygenerator.setGeometry(10, 70, 341, 41)
        self.keygenerator.setFont(QtGui.QFont('consola', 10))
        self.keygenerator.setAlignment(Qt.AlignLeft)
        self.keygenerator.setPlaceholderText('Type the of the Key')
        # self.keygenerator.setStyleSheet("background: #fff; color: #333; border: 1px solid #ccc;")

        self.btn_submit = QPushButton(self.frame2)
        self.btn_submit.setGeometry(130, 120, 121, 36)
        # self.btn_submit.setFont(QtGui.QFont('consola', 12))
        self.btn_submit.setObjectName("pushButton")
        self.btn_submit.setText("Create")
        self.btn_submit.setStyleSheet("background: #8be400; color: #fff; border: 4px solid transparent;")
        # self.btn_submit.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.btn_submit.clicked.connect(self.checkedme)

        self.plainText = QPlainTextEdit(self)
        self.plainText.setGeometry(10, 170, 341, 120)
        # self.plainText.setStyleSheet("background: #fff; border: 1px solid #ccc; font-size: 15px; color: #333;")
        self.plainText.setReadOnly(True)
        
        self.progressbar = QProgressBar(self)
        self.progressbar.setGeometry(10, 306, 341, 10)

        # self.labelInit = QLabel(self)
        # self.labelInit.setGeometry(190, 200, 81, 21)
        # self.labelInit.setText('InitTializing')

        # self.labelComplete = QLabel(self)
        # self.labelComplete.setGeometry(300, 200, 71, 16)
        # self.labelComplete.setText('completed')
    def checkedme(self):
        
        if self.radioButtonAsymmetric.isChecked():
            if self.keygenerator.text() != '':
                import shutil
                self.file_path = QtWidgets.QFileDialog.getExistingDirectory()

                for i in range(101):
                    time.sleep(0.02)
                    self.progressbar.setValue(i)
                key = RSA.generate(2048)
                self.bb = self.keygenerator.text()

                self.nameprivate = f'private_{self.bb}.key'
                self.namepublic = f'public_{self.bb}.key'

                
                private_key = key.export_key()
                with open(f'{self.nameprivate}', "wb") as file_out:

                    file_out.write(private_key)

                public_key = key.publickey().export_key()
                with open(f'{self.namepublic}', "wb") as file_out1:

                    file_out1.write(public_key)

                success = f'complete... \n {self.nameprivate} \n {self.namepublic} \n  On file Location {self.file_path} \n Generated successfully'

                self.plainText.setPlainText(success)

                self.basenm1 = os.path.join(os.getcwd(), f'{self.nameprivate}')
                self.basenm = os.path.join(os.getcwd(), f'{self.namepublic}')

                shutil.move(self.basenm1, self.file_path)
                shutil.move(self.basenm, self.file_path)
                self.keygenerator.clear()
            else:
                self.keygenerator.setText('Need a file name...')

        if self.radioButtonSymetric.isChecked():
            if self.keygenerator.text() != '':
                import shutil
                self.file_path = QtWidgets.QFileDialog.getExistingDirectory()
                print(self.file_path)

                for i in range(101):
                    time.sleep(0.02)
                    self.progressbar.setValue(i)
                # self.demo = LoadScreen()
                key = Fernet.generate_key()
                self.bb = self.keygenerator.text()

                with open(f'{self.bb}.key', 'wb') as mykey:
                    mykey.write(key)
                
                # print(self.bb)

                success = 'complete... \n key is generated successfully'

                self.plainText.setPlainText(success)

                self.basenm = os.path.join(os.getcwd(), f'{self.bb}.key')
                print(self.basenm)

                shutil.move(self.basenm, self.file_path)
                # print(self.basenm)
                # dropbox_path = "/Pythonfolder2"
                # dbx = dropbox.Dropbox(TOKEN)
                # dbx.files_upload(open(self.basenm, "rb").read(), dropbox_path)
                # print("[UPLOADED] {}".format(self.basenm))
            else:
                self.keygenerator.setText('Need a file name...')

    def createPublicKey(self):
        from Crypto.PublicKey import RSA

        key = RSA.generate(2048)
        private_key = key.export_key()
        file_out = open("private.key", "wb")
        file_out.write(private_key)
        file_out.close()

        public_key = key.publickey().export_key()
        file_out = open("public.key", "wb")
        file_out.write(public_key)
        file_out.close()


    def createEncryptkey1(self):
        
        if self.keygenerator.text() != '':
            import shutil
            self.file_path = QtWidgets.QFileDialog.getExistingDirectory()
            print(self.file_path)

            for i in range(101):
                time.sleep(0.02)
                self.progressbar.setValue(i)
            # self.demo = LoadScreen()
            key = Fernet.generate_key()
            self.bb = self.keygenerator.text()

            with open(f'{self.bb}.key', 'wb') as mykey:
                mykey.write(key)
            
            # print(self.bb)

            success = 'complete... \n key is generated successfully'

            self.plainText.setPlainText(success)

            self.basenm = os.path.join(os.getcwd(), f'{self.bb}.key')

            shutil.move(self.basenm, self.file_path)
        else:
            self.keygenerator.setText('Need a file name...')


class VerifyHashing(QWidget):
    def __init__(self):
        super().__init__()

        self.left =120
        self.top =200
        self.width =814
        self.height =530
        self.WindowAppVerify()
        # self.setStyleSheet("background: #f2f7d9; color:#333; ")
        self.setFixedSize(self.width, self.height)
        
        # self.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 rgba(234, 234, 234, 255), stop:0.366972 rgba(238, 238, 238, 255), stop:0.958716 rgba(215, 217, 218, 255));  font-size: 14px; color: #333; ")


    def WindowAppVerify(self):
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle(' App')

        self.boardframe = QFrame(self)
        self.boardframe.setGeometry(QtCore.QRect(-5, 60, 943, 450))
        self.boardframe.setFrameShape(QtWidgets.QFrame.Box)
        self.boardframe.setFrameShadow(QtWidgets.QFrame.Raised)
        self.boardframe.setObjectName("frame")
        # self.boardframe.setStyleSheet("background-color: #fff; color: #333;")

        self.frame1 = QFrame(self.boardframe)
        self.frame1.setGeometry(QtCore.QRect(450, 60, 362, 380))
        self.frame1.setFrameShape(QtWidgets.QFrame.Box)
        self.frame1.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame1.setObjectName("frame")
        # self.frame1.setStyleSheet("color: #fff; background: #f2f7d9;")

        self.label_Verify = QtWidgets.QLabel(self)
        self.label_Verify.setGeometry(QtCore.QRect(540, 20, 240, 21))
        self.label_Verify.setStyleSheet("font-size: 18px;\n"
                                   "font-weight: bold;\n"
                                   "background: none;\n"
                                   "border: none;\n"
                                   "color: #8be400;")
        self.label_Verify.setObjectName("label_3")
        self.label_Verify.setText('MESSAGE VERIFICATION')

        self.VerifyInputmessage = QLineEdit(self.frame1)
        self.VerifyInputmessage.setPlaceholderText('Your Message Here')
        self.VerifyInputmessage.setGeometry(10, 10, 270, 31)
        self.VerifyInputmessage.setFont(QtGui.QFont('consola', 10))
        self.VerifyInputmessage.setObjectName("qlineEdit")
        # self.VerifyInputmessage.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")


        self.btn_Vsearch = QPushButton(self.frame1)
        self.btn_Vsearch.setGeometry(285, 10, 71, 31)
        self.btn_Vsearch.setFont(QtGui.QFont('consola', 12))
        self.btn_Vsearch.setObjectName("pushButton")
        self.btn_Vsearch.setText("...")
        self.btn_Vsearch.setToolTip('Browse Message')
        # self.btn_Vsearch.setStyleSheet("background: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)); color: #eee; border-radius: 2px;")
        self.btn_Vsearch.clicked.connect(self.getBrowseVerifyMsg)
        self.btn_Vsearch.setStyleSheet("background: #8be400; color: #fff;")

        self.VerifyInput2key = QLineEdit(self.frame1)
        self.VerifyInput2key.setPlaceholderText('Your Key Here')
        self.VerifyInput2key.setGeometry(10, 50, 270, 31)
        self.VerifyInput2key.setFont(QtGui.QFont('consola', 10))
        # self.VerifyInput2key.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")

        self.btn_Vsearch2 = QPushButton(self.frame1)
        self.btn_Vsearch2.setGeometry(285, 50, 71, 31)
        self.btn_Vsearch2.setFont(QtGui.QFont('consola', 12))
        self.btn_Vsearch2.setObjectName("pushButton")
        self.btn_Vsearch2.setText("...")
        self.btn_Vsearch2.setToolTip('Browse Key')
        self.btn_Vsearch2.setStyleSheet("background: #8be400; color: #eee; border-radius: 2px;")
        self.btn_Vsearch2.clicked.connect(self.getBrowseVerifyKey)
        # self.btn_Vsearch2.setStyleSheet("background: #8be400; color: #fff;")


        self.VerifyInput2key3 = QLineEdit(self.frame1)
        self.VerifyInput2key3.setPlaceholderText('Your Key Here')
        self.VerifyInput2key3.setGeometry(10, 90, 270, 31)
        self.VerifyInput2key3.setFont(QtGui.QFont('consola', 10))
        # self.VerifyInput2key3.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")

        self.btn_Vsearch3 = QPushButton(self.frame1)
        self.btn_Vsearch3.setGeometry(285, 90, 71, 31)
        self.btn_Vsearch3.setFont(QtGui.QFont('consola', 12))
        self.btn_Vsearch3.setObjectName("pushButton")
        self.btn_Vsearch3.setText("...")
        self.btn_Vsearch3.setToolTip('Browse Key')
        self.btn_Vsearch3.setStyleSheet("background:  #8be400; color: #eee; border-radius: 2px;")
        self.btn_Vsearch3.clicked.connect(self.getBrowseVerifyKey)
        # self.btn_Vsearch3.setStyleSheet("background: #8be400; color: #fff;")

        self.btn_Verify = QPushButton(self.frame1)
        self.btn_Verify.setGeometry(120, 340, 121, 31)
        self.btn_Verify.setFont(QtGui.QFont('consola', 12))
        self.btn_Verify.setObjectName("pushButton")
        self.btn_Verify.setText("verify")
        self.btn_Verify.setToolTip('Verify message')
        self.btn_Verify.setStyleSheet("background: #8be400; color: #eee; border-radius: 2px;")
        self.btn_Verify.clicked.connect(self.Verify)
        # self.btn_Verify.setStyleSheet("background: #8be400; color: #fff;")

        self.messageVerifyCheck = QPlainTextEdit(self.frame1)
        self.messageVerifyCheck.setGeometry(5, 160, 351, 161)
        self.messageVerifyCheck.setFont(QtGui.QFont('consola', 10))
        # self.messageVerifyCheck.setStyleSheet("background: #fff; color: #333; border-radius: 4px;")


        self.frame2 = QFrame(self.boardframe)
        self.frame2.setGeometry(QtCore.QRect(10, 10, 361, 380))
        self.frame2.setFrameShape(QtWidgets.QFrame.Box)
        self.frame2.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame2.setObjectName("frame")
        # self.frame2.setStyleSheet("background: #f2f7d9; color: #fff;")

        self.label_signer = QtWidgets.QLabel(self)
        self.label_signer.setGeometry(QtCore.QRect(120, 20, 240, 21))
        self.label_signer.setStyleSheet("font-size: 18px;\n"
                                   "font-weight: bold;\n"
                                   "text-align: center;\n"
                                   "background: none;\n"
                                   "border: none;\n"
                                   "color: #8be400;")
        self.label_signer.setObjectName("label_3")
        self.label_signer.setText('MESSAGE SIGNATURE')

        self.SignInput = QLineEdit(self.frame2)
        self.SignInput.setPlaceholderText('Your Key Here')
        self.SignInput.setGeometry(5, 10, 271, 31)
        self.SignInput.setFont(QtGui.QFont('consola', 10))
        # self.SignInput.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")

        self.btn_Ssearch = QPushButton(self.frame2)
        self.btn_Ssearch.setGeometry(285, 10, 71, 31)
        self.btn_Ssearch.setFont(QtGui.QFont('consola', 12))
        self.btn_Ssearch.setObjectName("pushButton")
        self.btn_Ssearch.setText("...")
        self.btn_Ssearch.setToolTip('Search for Keys')
        self.btn_Ssearch.setStyleSheet("background: #8be400; color: #fff;")
        self.btn_Ssearch.clicked.connect(self.getBrowsermsg)

        self.SignInput2 = QLineEdit(self.frame2)
        self.SignInput2.setPlaceholderText('Your Key Here')
        self.SignInput2.setGeometry(5, 50, 271, 31)
        self.SignInput2.setFont(QtGui.QFont('consola', 10))
        # self.SignInput2.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")

        self.btn_Ssearch2 = QPushButton(self.frame2)
        self.btn_Ssearch2.setGeometry(285, 50, 71, 31)
        self.btn_Ssearch2.setFont(QtGui.QFont('consola', 12))
        self.btn_Ssearch2.setObjectName("pushButton")
        self.btn_Ssearch2.setText("...")
        self.btn_Ssearch2.setToolTip('Search for Keys')
        # self.btn_Ssearch2.setStyleSheet("background: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)); color: #eee; border-radius: 2px;")
        self.btn_Ssearch2.clicked.connect(self.getBrowsermsg)
        self.btn_Ssearch2.setStyleSheet("background: #8be400; color: #fff;")


        self.SignInput3 = QLineEdit(self.frame2)
        self.SignInput3.setPlaceholderText('Your Key Here')
        self.SignInput3.setGeometry(5, 90, 271, 31)
        self.SignInput3.setFont(QtGui.QFont('consola', 10))
        # self.SignInput3.setStyleSheet("background: #fff; border: none; color: #333; border-radius: 4px;")

        self.btn_Ssearch3 = QPushButton(self.frame2)
        self.btn_Ssearch3.setGeometry(285, 90, 71, 31)
        self.btn_Ssearch3.setFont(QtGui.QFont('consola', 12))
        self.btn_Ssearch3.setObjectName("pushButton")
        self.btn_Ssearch3.setText("...")
        self.btn_Ssearch3.setToolTip('Search for Keys')
        # self.btn_Ssearch3.setStyleSheet("background: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)); color: #eee; border-radius: 2px;")
        self.btn_Ssearch3.clicked.connect(self.getBrowsermsg)
        self.btn_Ssearch3.setStyleSheet("background: #8be400; color: #fff;")

        # self.btn_Ssearch = QPushButton(self.frame2)
        # self.btn_Ssearch.setGeometry(285, 20, 71, 31)
        # self.btn_Ssearch.setFont(QtGui.QFont('consola', 12))
        # self.btn_Ssearch.setObjectName("pushButton")
        # self.btn_Ssearch.setText("...")
        # self.btn_Ssearch.setToolTip('Search for Keys')
        # self.btn_Ssearch.setStyleSheet("background: #009999; color: #eee; border-radius: 2px;")
        # self.btn_Ssearch.clicked.connect(self.getBrowsermsg)

        self.btn_Sign = QPushButton(self.frame2)
        self.btn_Sign.setGeometry(120, 340, 121, 31)
        self.btn_Sign.setFont(QtGui.QFont('consola', 12))
        self.btn_Sign.setObjectName("pushButton")
        self.btn_Sign.setText("Sign")
        self.btn_Sign.setToolTip('Sign the Message')
        # self.btn_Sign.setStyleSheet("background: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)); color: #eee; border-radius: 2px;")
        self.btn_Sign.clicked.connect(self.sign)
        self.btn_Sign.setStyleSheet("background: #8be400; color: #fff;")

        self.messageWrite = QPlainTextEdit(self.frame2)
        self.messageWrite.setGeometry(5, 160, 351, 161)
        self.messageWrite.setFont(QtGui.QFont('consola', 10))
        # self.messageWrite.setStyleSheet("background: #fff; color: #333; border-radius: 4px;")

    def getBrowsermsg(self):
        try:

                self.file_path = QtWidgets.QFileDialog.getOpenFileName()
                self.basenm = os.path.dirname(self.file_path[0])
                self.basenamefile = os.path.basename(self.file_path[0])
                self.SignInput.setText(self.basenamefile)

        except FileNotFoundError:
                pass
    
    def getBrowseVerifyMsg(self):
        try:

                self.file_path = QtWidgets.QFileDialog.getOpenFileName()
                self.basenm = os.path.dirname(self.file_path[0])
                self.basenamefile = os.path.basename(self.file_path[0])
                self.VerifyInputmessage.setText(self.basenamefile)

        except FileNotFoundError:
                pass

    def getBrowseVerifyKey(self):
        try:

                self.file_path = QtWidgets.QFileDialog.getOpenFileName()
                self.basenm = os.path.dirname(self.file_path[0])
                self.basenamefile = os.path.basename(self.file_path[0])
                self.VerifyInput2key.setText(self.basenamefile)

        except FileNotFoundError:
                pass
    
    
    
    
    def sign(self):

        
        self.chat = self.messageWrite.toPlainText()

        with open('DepartmentMessage.txt', 'w') as kk:
            aa = kk.write(self.chat)

        with open('DepartmentMessage.txt', 'rb') as bb:
            c = bb.read()

        key = RSA.import_key(open(f'{self.SignInput.text()}').read())
        h = SHA256.new(c)

        '''
        once hashing is done, we need to create a sign object through 
        which we can sign a message
        '''

        signer=pkcs1_15.new(key)
        signature=signer.sign(h)
        #signature = pkcs1_15.new(key).sign(h)
        #signature_readable=codecs.getencoder('hex')(signature)
        print(signature.hex())


        with open("signature.pem", "wb") as file_out:
            file_out.write(signature)


    def Verify(self):
        key = RSA.import_key(open(f'{self.VerifyInput2key.text()}').read())

        with open(f'{self.VerifyInputmessage.text()}', "rb") as file_msg:

            message = file_msg.read()

        file_in = open("signature.pem", "rb")
        signature=file_in.read()
        file_in.close()
        h = SHA256.new(message)

        try:
            pkcs1_15.new(key).verify(h, signature)
            self.messageVerifyCheck.setPlainText("The signature is valid.")
        except (ValueError, TypeError):
            self.messageVerifyCheck.setPlainText("The signature is not valid.")
      

class LogForm(QWidget):
    def __init__(self):
        super().__init__()

        self.left =140
        self.top =200
        self.width =981
        self.height =500
        # self.setWindowFlag(QtCore.Qt.FramelessWindowHint)

        # self.setWindowModality(Qt.ApplicationModal)
        self.setFixedSize(self.width, self.height)
        
        self.setStyleSheet('background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)) ;'
                                        'color: #fff;'
                                        'border: none;'
                                        )

        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle('Login Form')

        

        self.frame2 = QFrame(self)
        self.frame2.setGeometry(QtCore.QRect(500, 0, 481, 500))
        self.frame2.setFrameShape(QtWidgets.QFrame.Box)
        self.frame2.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame2.setObjectName("frame")
        self.frame2.setStyleSheet("background: #eee; border: none;")

        self.image_frame = QtWidgets.QFrame(self.frame2)
        self.image_frame.setGeometry(QtCore.QRect(190, 20, 120, 130))
        self.image_frame.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.image_frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.image_frame.setObjectName("frame_10")
        self.image_frame.setStyleSheet("background: #fff;")

        
        self.label_user = QLabel(self.image_frame)
        self.label_user.setStyleSheet("border-image: url(:/image/user.ico); color: #000; border:none;")
        self.label_user.setGeometry(0, 0, 120, 130)

        self.label_user = QLabel(self)
        self.label_user.setStyleSheet("border-image: url(:/image/encrypted); color: #000; border:none;")
        self.label_user.setGeometry(110, 10, 250, 250)
        


        self.username = QLineEdit(self.frame2)
        self.username.setGeometry(98, 170, 300, 41)
        self.username.setFont(QtGui.QFont('consola', 10))
        self.username.setAlignment(Qt.AlignLeft)
        self.username.setPlaceholderText('Username')
        self.username.setToolTip('You are onusername Input')
        # self.username.textChanged[str].connect(self.login)
        self.username.setStyleSheet("background: #fff; font-size: 16px; color: #333; border-radius: 4px;")

        # self.frame_10 = QtWidgets.QFrame(self.frame2)
        # self.frame_10.setGeometry(QtCore.QRect(380, 170, 41, 41))
        # self.frame_10.setFrameShape(QtWidgets.QFrame.StyledPanel)
        # self.frame_10.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_10.setObjectName("frame_10")

        self.password = QLineEdit(self.frame2)
        self.password.setGeometry(98, 220, 300, 41)
        self.password.setFont(QtGui.QFont('consola', 10))
        self.password.setAlignment(Qt.AlignLeft)
        self.password.setPlaceholderText('Password')
        self.password.setToolTip('You are password Input')
        # self.password.setValidator(QIntValidator())
        # self.password.textChanged[str].connect(self.login)
        self.password.setEchoMode(QLineEdit.Password)
        self.password.setStyleSheet("background: #fff; font-size: 16px; color: #333; border-radius: 4px;")

        # self.frame_1b = QtWidgets.QFrame(self.frame2)
        # self.frame_1b.setGeometry(QtCore.QRect(380, 220, 41, 41))
        # self.frame_1b.setFrameShape(QtWidgets.QFrame.StyledPanel)
        # self.frame_1b.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_1b.setObjectName("frame_10")

        self.btn_submit = QPushButton(self.frame2)
        self.btn_submit.setGeometry(200, 280, 111, 41)
        self.btn_submit.setFont(QtGui.QFont('consola', 12))
        self.btn_submit.setObjectName("pushButton")
        self.btn_submit.setText("Login")
        self.btn_submit.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.btn_submit.setStyleSheet('QPushButton'
                                        '{'
                                        'background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0.54171, y2:0, stop:0 rgba(35, 160, 80, 255), stop:1 rgba(76, 201, 135, 255)) ;'
                                        'color: #fff;'
                                        'border: none;'
                                        '}'
                                        'QPushButton#pushButton:pressed'
                                        '{'
                                            'padding-left: 5px;'
                                            'padding-top: 5px;'
                                            '}'
                                                )
        
        self.btn_submit.clicked.connect(self.login)
        # self.btn_submit.clicked.connect(self.check)


        self.progressbar = QProgressBar(self.frame2)
        self.progressbar.setGeometry(40, 330, 401, 10)
        


        


        self.plainText = QPlainTextEdit(self.frame2)
        self.plainText.setGeometry(20, 360, 441, 131)
        self.plainText.setStyleSheet("background: #fff; border: 1px solid #ccc; border-radius: 4px; font-size: 15px; color: #333;")
        self.plainText.setReadOnly(True)

        # self.frame_15 = QtWidgets.QFrame(self)
        # self.frame_15.setGeometry(QtCore.QRect(100, 20, 120, 120))
        # self.frame_15.setStyleSheet("border-image: url(:/image/371-3715555_nasarawa-state-university-keffi-nsuk.png);")
        # self.frame_15.setFrameShape(QtWidgets.QFrame.StyledPanel)
        # self.frame_15.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_15.setObjectName("frame_15")

        # self.frame_16 = QtWidgets.QFrame(self)
        # self.frame_16.setGeometry(QtCore.QRect(110, 10, 250, 250))
        # self.frame_16.setStyleSheet("border-image: url(:/image/PngItem_2666090.png);")
        # self.frame_16.setFrameShape(QtWidgets.QFrame.StyledPanel)
        # self.frame_16.setFrameShadow(QtWidgets.QFrame.Raised)
        # self.frame_16.setObjectName("frame_15")
        
        self.Logoname = QLabel(self)
        self.Logoname.setGeometry(35, 310, 440, 90)
        self.Logoname.setText('DATA ENCRYPTION AND DECRYPTION SECURITY SYSTEM')
        self.Logoname.setStyleSheet('background: none; font-weight: bold; border:none; color: #FFF; font-size: 15px;')

        self.Logoname2 = QLabel(self)
        self.Logoname2.setFont(QtGui.QFont('weight: 25'))
        self.Logoname2.setGeometry(100, 260, 420, 80)
        self.Logoname2.setText('SECURE DATA')
        self.Logoname2.setStyleSheet('background: none; color: #eee; border:none; font-size: 30px; font-weight: bold;')

        self.label_animation = QLabel(self)
        self.label_animation.setStyleSheet("background: none; color: #fff; border:none;")
        self.movie = QMovie('64x64.gif')
        self.label_animation.setGeometry(200, 380, 118, 100)
        self.label_animation.setMovie(self.movie)

        timer = QTimer()
        self.startAnimation()
        # timer.singleShot(8000)

        self.show()

    def startAnimation(self):
        self.movie.start()


    def login(self):
        if self.username.text() == 'joe' and self.password.text() == '12345':

            for i in range(101):
                time.sleep(0.03)
                self.progressbar.setValue(i)
            
            self.close()

            self.MainWindow = QtWidgets.QMainWindow()
            self.ui = Ui_MainWindow()
            self.ui.setupUi(self.MainWindow)
            self.MainWindow.show()

            
        else:
            name = socket.gethostname()
            self.plainText.setPlainText(f'{name}: trying to login \n Not correct')
            


class ChatingApp(QWidget):
    def __init__(self):
        super().__init__()

        self.left =100
        self.top =400
        self.width =464
        self.height =540
        self.ChatApp()
        # self.setWindowModality(Qt.ApplicationModal)
        self.setFixedSize(self.width, self.height)
        # self.setStyleSheet("background-color: #fff;")


    def ChatApp(self):
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle('Chat App')

        self.frame = QFrame(self)
        self.frame.setGeometry(QtCore.QRect(-5, 0, 581, 300))
        self.frame.setFrameShape(QtWidgets.QFrame.Box)
        self.frame.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame.setObjectName("frame")
        # self.frame.setStyleSheet("color: #fff; background: #f2f7d9")

        # self.labelPhone = QLabel(self.frame)
        # self.labelPhone.setGeometry(30, 10, 81, 21)
        # self.labelPhone.setText('Encryption Key')
        # self.labelPhone.setStyleSheet('color: #eee')

        self.phone_number = QLineEdit(self.frame)
        self.phone_number.setPlaceholderText('Your Phone Number')
        self.phone_number.setGeometry(30, 30, 281, 41)
        self.phone_number.setFont(QtGui.QFont('consola', 10))
        # self.phone_number.setStyleSheet("background: #fff; border: 1px solid #ccc; color: #333; border-radius: 4px;")

        # self.labelInit = QLabel(self.frame)
        # self.labelInit.setGeometry(30, 75, 81, 21)
        # self.labelInit.setText('Encryption Key')
        # self.labelInit.setStyleSheet('color: #eee')

        self.key = QLineEdit(self.frame)
        self.key.setPlaceholderText('Your Encrypt/Decrypt Key Here')
        self.key.setGeometry(30, 100, 380, 41)
        self.key.setFont(QtGui.QFont('consola', 10))
        # self.key.setStyleSheet("background: #fff; border: 1px solid #ccc; color: #333; border-radius: 4px;")


        self.btn_browse_key = QPushButton(self)
        self.btn_browse_key.setGeometry(410, 100, 41, 41)
        self.btn_browse_key.setFont(QtGui.QFont('consola', 12))
        self.btn_browse_key.setText('...')
        self.btn_browse_key.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.btn_browse_key.setToolTip('Browse a File')
        self.btn_browse_key.clicked.connect(self.GetKeyBrowser)
        self.btn_browse_key.setStyleSheet("background: #8be400; color: #fff")


        self.labelComplete = QLabel(self)
        self.labelComplete.setGeometry(200, 160, 71, 16)
        self.labelComplete.setText('Chat Now')
        self.labelComplete.setStyleSheet('background: none; color: #eee; border: none')

        self.messages = QPlainTextEdit(self)
        self.messages.setGeometry(10, 175, 446, 171)
        self.messages.setFont(QtGui.QFont('consola', 10))
        # self.messages.setStyleSheet("background: #f8fbfc; color: #4681b4;; border-radius: 4px;")

        self.btn_encrypt = QPushButton(self)
        self.btn_encrypt.setGeometry(100, 355, 42, 42)
        self.btn_encrypt.setFont(QtGui.QFont('consola', 12))
        self.btn_encrypt.setObjectName("pushButton")
        # self.btn_encrypt.setText("Encrypt")
        self.btn_encrypt.setToolTip('Click to Encrypt')
        self.btn_encrypt.clicked.connect(self.chatEncrypt)
        # self.btn_encrypt.setStyleSheet('background: #002266; color: #eee; border-radius:8px')
        self.btn_encrypt.setStyleSheet('QPushButton'
            '{'
            'border-image: url(:/image/icons/lock_32x32.png);'
            'color: #333;'
             '}'
             'QPushButton::Hover'
            '{'
            'background: #181825;'
            '}')

        self.btn_decrypt = QPushButton(self)
        self.btn_decrypt.setGeometry(300, 355, 42, 42)
        self.btn_decrypt.setFont(QtGui.QFont('consola', 13))
        self.btn_decrypt.setObjectName("pushButton")
        self.btn_decrypt.setToolTip('Click to Decrypt')
        self.btn_decrypt.clicked.connect(self.chatDecrypt)
        # self.btn_decrypt.setStyleSheet('background: #00b3a1; color: #eee; border-radius:8px')
        self.btn_decrypt.setStyleSheet('QPushButton'
            '{'
            'border-image: url(:/image/icons/unlock_32x32.png);'
             '}'
             'QPushButton::Hover'
            '{'
            'background: #181825;'
            '}')
        
        self.btn_send = QPushButton(self)
        self.btn_send.setGeometry(120, 430, 201, 41)
        self.btn_send.setFont(QtGui.QFont('consola', 13))
        self.btn_send.setObjectName("pushButton")
        self.btn_send.setText("Send")
        self.btn_send.setToolTip('send Text')
        self.btn_send.clicked.connect(self.sendMessage)
        self.btn_send.setStyleSheet('QPushButton'
                                    '{'
                                    'background: #8be400;'
                                    'color: #eee; border-radius:4px'
                                    '}'
                                    'QPushButton::Hover'
                                    '{'
                                    'background: #8be400;'
                                    '}')


    def GetKeyBrowser(self):
        try:
                self.file_path = QtWidgets.QFileDialog.getOpenFileName()
                self.basenm = os.path.dirname(self.file_path[0])
                self.basenamefile = os.path.basename(self.file_path[0])
                self.bb = f'{os.path.join(self.basenm, self.basenamefile)}'
                # print(self.path_lineEdit.text(), 'hello')
                with open(f'{self.bb}', 'rt') as mykey2:
                        key2 = mykey2.read()
                        self.key.setText(str(key2))

                with open('text.txt', 'rt') as original_files:
                    original = original_files.read()
                    self.messages.setPlainText(str(original))

        except FileNotFoundError:
                pass


    def chatEncrypt(self):
        day = datetime.now()
        day2 = day.strftime("%A")
        #............................................
        now = datetime.now()
        mtime = now.strftime("%I: %M: %p")
        #................................................
        today = datetime.today()
        STAMP_DATE = today.strftime("%d/%m/%y")

        print(STAMP_DATE, mtime)

        self.ph = self.phone_number.text()
        self.readkey = self.key.text()


        f = Fernet(self.readkey)

        self.chat = self.messages.toPlainText()

        with open('mainMessage.txt', 'w') as kk:
            aa = kk.write(self.chat)

        with open('mainMessage.txt', 'rb') as bb:
            c = bb.read()

        encrypted = f.encrypt(c)
        
        with open('mainMessage.txt', 'wb') as bb2:
            c2 = bb2.write(encrypted)
        
        with open('mainMessage.txt', 'rt') as mm:
            mb = mm.read()
        self.messages.setPlainText(str(mb))


    def sendMessage(self):
            try:
                import pywhatkit as kit
                # print(self.messages.toPlainText())
                kit.sendwhatmsg(self.phone_number.text(), self.messages.toPlainText(), 17, 10)

                print('i have send message')
            except:
                QMessageBox.warning(self, "MESSAGE", "Check your internet connection")


        
        

        ###########################################################
        # self.chat = self.messages.toPlainText()

        # self.encode_message = self.chat.encode("ascii")
        # self.base64_byte = base64.b64encode(self.encode_message)
        # self.encrypt = self.base64_byte.decode("ascii")
        # self.messages.setPlainText(str(self.encrypt))
        ###########################################################


            

    def chatDecrypt(self):

        day = datetime.now()
        day2 = day.strftime("%A")
        #............................................
        now = datetime.now()
        mtime = now.strftime("%I: %M: %p")
        #................................................
        today = datetime.today()
        STAMP_DATE = today.strftime("%d/%m/%y")

        print(STAMP_DATE, mtime)


        self.ph = self.phone_number.text()
        self.readkey = self.key.text()

        f = Fernet(self.readkey)

        with open('mainMessage.txt', 'rb') as bb:
            c = bb.read()

        decrypted = f.decrypt(c)
        
        with open('mainMessage.txt', 'wb') as bb2:
            c2 = bb2.write(decrypted)
        
        with open('mainMessage.txt', 'rt') as mm:
            mb = mm.read()
        self.messages.setPlainText(str(mb))
    

        #######################################################
        # self.chat = self.messages.toPlainText()

        # self.encode_message = self.chat.encode("ascii")
        # self.base64_byte = base64.b64decode(self.encode_message)
        # self.encrypt = self.base64_byte.decode("ascii")
        # self.messages.setPlainText(str(self.encrypt))
        ###########################################################


            
    # def Chat(self):

    #     pwt.sendwhatmsg(f'{self.phone_number}, {self.messages}, 20, 30')
        

class Encryption_Decrypt_Form(QMainWindow):
    def __init__(self, parent):
        super().__init__(parent)

        # self.setStyleSheet("background-color: qlineargradient(spread:pad, x1:0, y1:0, x2:0, y2:1, stop:0 rgba(234, 234, 234, 255), stop:0.366972 rgba(238, 238, 238, 255), stop:0.958716 rgba(215, 217, 218, 255)); color: #333")
        # self.setStyleSheet("color: #fff;")
        # self.setWindowModality(Qt.ApplicationModal)
        self.resize(576, 451)

        self.frame3 = QFrame(self)
        self.frame3.setGeometry(QtCore.QRect(2, 0, 580, 240))
        self.frame3.setFrameShape(QtWidgets.QFrame.StyledPanel)
        self.frame3.setFrameShadow(QtWidgets.QFrame.Raised)
        self.frame3.setObjectName("frame")
        # self.frame3.setStyleSheet("background-color: #f2f7d9; color:#333;")


        self.keygenerator = QLineEdit(self.frame3)
        self.keygenerator.setGeometry(40, 60, 481, 41)
        self.keygenerator.setFont(QtGui.QFont('consola', 10))
        # self.keygenerator.setStyleSheet("background: #fff; border: 1px solid #ccc; color: #333;")
        self.keygenerator.setPlaceholderText('insert Your Single key')

        self.btn_browse = QPushButton(self.frame3)
        self.btn_browse.setGeometry(525, 60, 41, 41)
        self.btn_browse.setFont(QtGui.QFont('consola', 12))
        self.btn_browse.setText('...')
        self.btn_browse.setCursor(QtGui.QCursor(Qt.PointingHandCursor))
        self.btn_browse.setToolTip('Browse a File')
        self.btn_browse.clicked.connect(self.getBrowser)
        self.btn_browse.setStyleSheet("background: #8be400; color: #fff;")


        self.path_lineEdit = QLineEdit(self.frame3)
        self.path_lineEdit.setGeometry(10, 10, 552, 41)
        self.path_lineEdit.setFont(QtGui.QFont('consola', 10))
        # self.path_lineEdit.setText(signal)
        # self.path_lineEdit.setStyleSheet("background: #fff; border: 1px solid #ccc; color: #333;")


        self.btn_encrypt = QPushButton(self.frame3)
        self.btn_encrypt.setGeometry(165, 130, 42, 42)
        self.btn_encrypt.setFont(QtGui.QFont('consola', 13))
        self.btn_encrypt.setObjectName("pushButton")
        # self.btn_encrypt.setText("Encrypt")
        self.btn_encrypt.setToolTip('Click to Encrypt')
        self.btn_encrypt.clicked.connect(self.EncryptSource)
        self.btn_encrypt.setStyleSheet('background: #8be400; color: #fff;')
        self.btn_encrypt.setStyleSheet('QPushButton'
            '{'
            'border-image: url(:/image/icons/lock_32x32.png);'
            'padding-left: 5px;'
            'padding-top: 5px;'
            # 'background: #8be400;'
             '}'
             'QPushButton::Hover'
            '{'
            'background: #ddd;'
            '}')
            # 'QPushButton#pushButton:pressed'
            # '{'
            #     'padding-left: 5px;'
            #     'padding-top: 5px;'
            #     'background: #8be400;'
            # '}')


        self.btn_decrypt = QPushButton(self.frame3)
        self.btn_decrypt.setGeometry(360, 132, 42, 42)
        self.btn_decrypt.setFont(QtGui.QFont('consola', 13))
        self.btn_decrypt.setObjectName("pushButton")
        # self.btn_decrypt.setText("Decrypt")
        self.btn_decrypt.setToolTip('Click to Decrypt')
        self.btn_decrypt.clicked.connect(self.DecryptSource)
        # self.btn_decrypt.setStyleSheet('background: #002266; color: #eee; border-radius:8px')
        self.btn_decrypt.setStyleSheet('QPushButton'
            '{'
            'border-image: url(:/image/icons/unlock_32x32.png);'
             '}'
             'QPushButton::Hover'
            '{'
            'background: #ddd;'
            '}')

        self.progressbar = QProgressBar(self)
        self.progressbar.setGeometry(10, 230, 561, 10)
        

        self.plainText = QPlainTextEdit(self)
        self.plainText.setGeometry(40, 245, 500, 190)
        # self.plainText.setStyleSheet("background: #f8fbfc; border: 1px solid #ccc;  color: #333;")

        self.labelEncrypt = QLabel(self.frame3)
        self.labelEncrypt.setGeometry(170, 190, 81, 21)
        self.labelEncrypt.setText('Encrypt')

        self.labelDecrypt = QLabel(self.frame3)
        self.labelDecrypt.setGeometry(360, 190, 71, 16)
        self.labelDecrypt.setText('Decrypt')

    def getBrowser(self):
        try:

                self.file_path = QtWidgets.QFileDialog.getOpenFileName()
                self.basenm = os.path.dirname(self.file_path[0])
                self.basenamefile = os.path.basename(self.file_path[0])

                self.bbkey = f'{os.path.join(self.basenm, self.basenamefile)}'
                with open(f'{self.bbkey}', 'rt') as mykey2:
                        key2 = mykey2.read()
                        self.keygenerator.setStyleSheet("background: #fff; color: #333")
                        self.keygenerator.setText(str(key2))
        except FileNotFoundError:
                pass
    

    def EncryptSource(self):
        day = datetime.now()
        day2 = day.strftime("%A")
        #............................................
        now = datetime.now()
        mtime = now.strftime("%I: %M: %p")
        #................................................
        today = datetime.today()
        STAMP_DATE = today.strftime("%d/%m/%y")

        conn = sqlite3.connect('encryption_decryption_storebase.db')
        # print(' database open successfully')
        
        try:
            for i in range(101):
                time.sleep(0.02)
                self.progressbar.setValue(i)

            self.dd = self.path_lineEdit.text()
            # self.basenamefile = os.path.basename(self.dd)
            self.encryptkeyread = self.keygenerator.text()
            if self.keygenerator.text() == '':
                self.keygenerator.setText('key is needed')
            
            else:

                f = Fernet(self.encryptkeyread)

                with open(self.dd, 'rb') as original_files:
                    original = original_files.read()

                    encrypted = f.encrypt(original)

                with open(self.dd, 'wb') as encrypted_files:
                    encrypted_files.write(encrypted)

                success = 'Encrypted SuccessFully'
                self.plainText.setPlainText(f'{success}')

                insert_data = "INSERT INTO encryption_decryption_storebase (key_name, key_value, filename, stamp_date, time, day) VALUES('"+self.basenamefile+"', '" + self.encryptkeyread + "', '"+self.path_lineEdit.text()+"', '" + STAMP_DATE+ "', '" +mtime+ "', '" +day2 + "') "
                conn.execute(insert_data)
                conn.commit()
                # print('Record created successfully')
                conn.close()
        except AttributeError:
            self.message = "Oops Key is needed...."
            self.keygenerator.setStyleSheet("color: red")
            self.keygenerator.setText(self.message)


    def DecryptSource(self):

        try:
            for i in range(101):
                time.sleep(0.03)
                self.progressbar.setValue(i)

            self.bb = self.path_lineEdit.text()
            self.readkey = self.keygenerator.text()


            fd = Fernet(self.readkey)

            with open(self.bb, 'rb') as encrypted_files:
                encryted = encrypted_files.read()

            decrypt = fd.decrypt(encryted)

            with open(self.bb, 'wb') as decrypted_files:
                decrypted_files.write(decrypt)

            # with open(self.bb, 'rt') as mykey3:
            #     key3 = mykey3.read()
            success = 'Decrypted SuccessFully'
            self.plainText.setPlainText(f'{success}')
        except AttributeError:
            self.message = "Oops Key is needed...."
            self.keygenerator.setStyleSheet("color: red")
            self.keygenerator.setText(self.message)




class DatabaseForm(QWidget):
    def __init__(self):
        super().__init__()

        self.left =60
        self.top =100
        self.width =1200
        self.height =600
        self.DatabaseWindow()
        # self.setStyleSheet("background: #f2f7d9; color: #333")
        self.resize(self.width, self.height)
        # self.setWindowModality(Qt.ApplicationModal)
        
        # self.setStyleSheet("background: #009b77; color: #4681b4;; border: 1px solid #3477b6;")
        self.loadTable()

    def DatabaseWindow(self):
        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle(' Database Table')

        # self.btn_database = QPushButton(self)
        # self.btn_database.setGeometry(990, 5, 101, 31)
        # # self.btn_database.setFont(QtGui.QFont('consola', 12))
        # self.btn_database.setObjectName("pushButton")
        # self.btn_database.setText("Query")
        # self.btn_database.setToolTip('Query the Database')
        # self.btn_database.setStyleSheet("background: #eee; color: #333; border-radius: 2px;")
        # self.btn_database.clicked.connect(self.loadTable)
        

      
        self.tableWidget = QtWidgets.QTableWidget(self)
        self.tableWidget.setGeometry(QtCore.QRect(0, 5, 1230, 1500))
        self.tableWidget.setFrameShape(QFrame.WinPanel)
        self.tableWidget.setFont(QtGui.QFont('consola', 10))
        self.tableWidget.setGridStyle(QtCore.Qt.DotLine)
        self.tableWidget.setRowCount(1)
        self.tableWidget.setColumnCount(7)
        header = ['Id','KeyName', 'KeyValue', 'FileName', 'Date', 'Time', 'Day']
        self.tableWidget.setHorizontalHeaderLabels(header)
        self.tableWidget.setColumnWidth(0, 20)
        self.tableWidget.setColumnWidth(1, 200)
        self.tableWidget.setColumnWidth(2, 245)
        self.tableWidget.setColumnWidth(3, 245)
        self.tableWidget.setColumnWidth(4, 140)
        self.tableWidget.setColumnWidth(5, 140)
        self.tableWidget.setColumnWidth(6, 140)
        
        self.tableWidget.verticalHeader().setCascadingSectionResizes(True)
        # self.tableWidget.setStyleSheet("background: #fff; color: #333; border: 1px solid #ccc; font-size: 13px;")


    def loadTable(self):
        self.conn = sqlite3.connect('encryption_decryption_storebase.db')
        self.cursor="SELECT * FROM encryption_decryption_storebase"
        self.result = self.conn.execute(self.cursor)
        self.tableWidget.setRowCount(0)

        for row_num, row_data in enumerate(self.result):
            self.tableWidget.insertRow(row_num)
            for column_num, data in enumerate(row_data):
                self.tableWidget.setItem(row_num, column_num, QtWidgets.QTableWidgetItem(str(data)))
        
        self.conn.commit()
        self.conn.rollback()



class EmailForm(QWidget):
    def __init__(self):
        super().__init__()

        self.left =140
        self.top =600
        self.width =400
        self.height =600
        # self.setWindowFlag(QtCore.Qt.FramelessWindowHint)

        # self.setWindowModality(Qt.ApplicationModal)
        self.setFixedSize(self.width, self.height)
        
        # self.setStyleSheet("background: #fff;")

        self.setGeometry(self.top, self.left, self.width, self.height)
        self.setWindowTitle('Email Form')

        self.label = QtWidgets.QLabel(self)
        self.label.setGeometry(QtCore.QRect(10, 10, 100, 31))
        font = QtGui.QFont()
        font.setPointSize(21)
        font.setBold(True)
        font.setWeight(20)
        self.label.setFont(font)
        self.label.setObjectName("label")

        self.label_2 = QtWidgets.QLabel(self)
        self.label_2.setGeometry(QtCore.QRect(10, 40, 231, 17))
        font = QtGui.QFont()
        font.setPointSize(10)
        font.setBold(True)
        font.setWeight(75)
        self.label_2.setFont(font)
        self.label_2.setObjectName("label_2")
        
        self.groupBox = QtWidgets.QGroupBox(self)
        self.groupBox.setGeometry(QtCore.QRect(5, 70, 390, 500))
        # self.groupBox.setStyleSheet('background: #f2f7d9')
        self.groupBox.setObjectName("groupBox")

        self.email = QtWidgets.QLabel(self.groupBox)
        self.email.setGeometry(QtCore.QRect(10, 26, 71, 21))
        self.email.setFont(font)
        self.email.setObjectName("email")
        self.email.setText("Email")
        
        self.username = QtWidgets.QLineEdit(self.groupBox)
        self.username.setGeometry(QtCore.QRect(10, 45, 311, 36))
        self.username.setObjectName("username")
        self.username.setFont(QtGui.QFont('Arial', 13))
        # self.username.setStyleSheet('background: #fff')

        self.passwordlbl = QtWidgets.QLabel(self.groupBox)
        self.passwordlbl.setGeometry(QtCore.QRect(10, 85, 71, 21))
        self.passwordlbl.setFont(font)
        self.passwordlbl.setObjectName("passwordlbl")
        self.passwordlbl.setText("Password")

        self.password = QtWidgets.QLineEdit(self.groupBox)
        self.password.setGeometry(QtCore.QRect(10, 105, 311, 36))
        self.password.setObjectName("password")
        self.password.setFont(QtGui.QFont('Arial', 13))
        # self.password.setStyleSheet('background: #fff')

        #################################################
        self.reciever = QtWidgets.QLabel(self.groupBox)
        self.reciever.setGeometry(QtCore.QRect(10, 145, 71, 21))
        self.reciever.setFont(font)
        self.reciever.setObjectName("reciever")
        self.reciever.setText("To")
        
        self.recieverMail = QtWidgets.QLineEdit(self.groupBox)
        self.recieverMail.setGeometry(QtCore.QRect(10, 166, 311, 36))
        self.recieverMail.setObjectName("recieverMail")
        self.recieverMail.setFont(QtGui.QFont('Arial', 13))
        # self.recieverMail.setStyleSheet('background: #fff')

        self.subjectlbl = QtWidgets.QLabel(self.groupBox)
        self.subjectlbl.setGeometry(QtCore.QRect(10, 208, 71, 21))
        self.subjectlbl.setFont(font)
        self.subjectlbl.setObjectName("subjectlbl")
        self.subjectlbl.setText("Subject")

        self.subject = QtWidgets.QLineEdit(self.groupBox)
        self.subject.setGeometry(QtCore.QRect(10, 228, 311, 36))
        self.subject.setObjectName("subject")
        self.subject.setFont(QtGui.QFont('Arial', 13))
        # self.subject.setStyleSheet('background: #fff')

        self.bodylbl = QtWidgets.QLabel(self.groupBox)
        self.bodylbl.setGeometry(QtCore.QRect(10, 275, 71, 21))
        self.bodylbl.setFont(font)
        self.bodylbl.setObjectName("bodylbl")
        self.bodylbl.setText("Body")

        self.body = QtWidgets.QPlainTextEdit(self.groupBox)
        self.body.setGeometry(QtCore.QRect(10, 295, 311, 100))
        self.body.setObjectName("body")
        self.body.setFont(QtGui.QFont('Arial', 13))
        # self.body.setStyleSheet('background: #fff; border: 1px solid #ccc')

        self.pushButton = QtWidgets.QPushButton(self.groupBox)
        self.pushButton.setGeometry(QtCore.QRect(70, 430, 201, 36))
        self.pushButton.setObjectName("pushButton")
        self.pushButton.setStyleSheet('background: #8be400; color: #fff')
        # self.pushButton.clicked.connect(self.closeReg)
        
        self.progressBar = QtWidgets.QProgressBar(self)
        self.progressBar.setGeometry(QtCore.QRect(20, 570, 330, 8))
        self.progressBar.setProperty("value", 24)
        self.progressBar.setObjectName("progressBar")

        self.pushButton.setText("Submit")
        self.label.setText("Custom Email")
        
        self.label_2.setText("Using the form below to send Email")
        # self.groupBox.setTitle("Form")
    # def save(self):
    #     f = open("passfile.txt", "a" )
    #     f.write(f'\n {self.lineEdit_2.text()}') 

    #     self.readdata()

    # def readdata(self):
    #     fd = open("passfile.txt", "r" )
    #     m = fd.read()
    #     self.plainTextEdit.setPlainText(str(m))

    # def closeReg(self):
    #     if self.username.text() == 'joe' and self.password.text() == '123':
            

    #         for i in range(50):
    #             time.sleep(0.03)
    #             self.progressBar.setValue(i)
    #         self.close()

    #         window = Attendance()
    #         window.username.setText(f'Welcome {self.username.text()}')
    #         window.show()

    #         window22 = Window_Form()
    #         window22.hide()



if __name__ == "__main__":
    import sys
    app = QtWidgets.QApplication(sys.argv)
    app.setStyle('Fusion')
    # app.setStyleSheet(qdarkstyle.load_stylesheet())
    # app.setStyleSheet(qdarktheme.load_stylesheet("light"))
    # app.setStyleSheet(qdarkgraystyle.load_stylesheet())
    MainWindow = QtWidgets.QMainWindow()

    ui = Ui_MainWindow()
    ui.setupUi(MainWindow)
    MainWindow.show()

    # win = LogForm()
    # win.show()
    # win = ChatingApp()
    # win.show()

    sys.exit(app.exec_())